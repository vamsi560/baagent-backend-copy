#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to process Personal Auto documents and add them to pgvector database
"""

import os
import sys

# Fix Windows console encoding
if sys.platform == 'win32':
    try:
        import codecs
        if hasattr(sys.stdout, 'buffer'):
            sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
        if hasattr(sys.stderr, 'buffer'):
            sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')
    except:
        pass  # Continue without encoding fix if it fails

from pathlib import Path
from docx import Document as DocxDocument
import uuid
from database import add_to_vector_db, embedding_model, PINECONE_ENABLED
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Personal Auto documents folder
PERSONAL_AUTO_FOLDER = Path("personalauto")

def extract_text_from_docx(file_path):
    """Extract text from DOCX file"""
    try:
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"❌ Error extracting text from {file_path}: {e}")
        return None

def chunk_text(text, chunk_size=1000, chunk_overlap=200):
    """Split text into chunks for better vector search"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    return splitter.split_text(text)

def process_personal_auto_documents():
    """Process all Personal Auto documents and add to Pinecone vector database"""
    if not PINECONE_ENABLED:
        print("ERROR: Pinecone is not enabled. Please enable it in database.py")
        return False
    
    if not embedding_model:
        print("❌ Embedding model not initialized")
        return False
    
    if not PERSONAL_AUTO_FOLDER.exists():
        print(f"❌ Personal Auto folder not found: {PERSONAL_AUTO_FOLDER}")
        return False
    
    print("Processing Personal Auto documents...")
    print(f"Folder: {PERSONAL_AUTO_FOLDER}")
    
    # Get all DOCX files
    docx_files = list(PERSONAL_AUTO_FOLDER.glob("*.docx"))
    
    if not docx_files:
        print(f"WARNING: No DOCX files found in {PERSONAL_AUTO_FOLDER}")
        return False
    
    print(f"SUCCESS: Found {len(docx_files)} documents to process")
    
    total_chunks = 0
    
    for docx_file in docx_files:
        print(f"\nProcessing: {docx_file.name}")
        
        # Extract text
        text = extract_text_from_docx(docx_file)
        if not text:
            print(f"WARNING: Skipping {docx_file.name} - no text extracted")
            continue
        
        print(f"   Extracted {len(text)} characters")
        
        # Split into chunks
        chunks = chunk_text(text)
        print(f"   Split into {len(chunks)} chunks")
        
        # Add each chunk to vector database
        for i, chunk in enumerate(chunks):
            chunk_id = str(uuid.uuid4())
            meta = {
                'id': chunk_id,
                'document_name': docx_file.name,
                'document_path': str(docx_file),
                'chunk_index': i,
                'total_chunks': len(chunks),
                'lob': 'personal_auto',
                'source': 'training_data'
            }
            
            add_to_vector_db(
                content=chunk,
                meta=meta,
                lob='personal_auto'
            )
            total_chunks += 1
        
        print(f"   SUCCESS: Added {len(chunks)} chunks to vector database")
    
    print(f"\nSUCCESS: Processing complete!")
    print(f"   Total documents processed: {len(docx_files)}")
    print(f"   Total chunks added: {total_chunks}")
    
    return True

if __name__ == "__main__":
    print("=" * 60)
    print("Personal Auto Documents Processing Script")
    print("=" * 60)
    
    success = process_personal_auto_documents()
    
    if success:
        print("\nSUCCESS: All documents processed successfully!")
        sys.exit(0)
    else:
        print("\nERROR: Processing failed!")
        sys.exit(1)

